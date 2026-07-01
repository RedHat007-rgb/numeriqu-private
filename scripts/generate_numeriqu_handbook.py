from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = "/Users/basanireddy/Desktop/test-1234/Numeriqu_Engineering_Onboarding_Handbook.docx"


def set_page_background(section):
    sect_pr = section._sectPr
    pg_borders = sect_pr.first_child_found_in("w:pgBorders")
    if pg_borders is None:
        pg_borders = OxmlElement("w:pgBorders")
        pg_borders.set(qn("w:offsetFrom"), "page")
        sect_pr.append(pg_borders)
    for edge in ("top", "left", "bottom", "right"):
        el = pg_borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            pg_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:space"), "24")
        el.set(qn("w:color"), "D9DEE8")


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def configure_styles(document):
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in (
        ("Title", 24, RGBColor(20, 33, 61)),
        ("Heading 1", 16, RGBColor(20, 33, 61)),
        ("Heading 2", 13, RGBColor(36, 71, 120)),
        ("Heading 3", 11, RGBColor(60, 60, 60)),
    ):
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    if "Callout" not in styles:
        callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
        callout.font.name = "Aptos"
        callout.font.size = Pt(10.5)
        callout.font.italic = True
        callout.font.color.rgb = RGBColor(52, 73, 94)
        callout.paragraph_format.left_indent = Inches(0.2)
        callout.paragraph_format.space_after = Pt(8)


def add_cover(document):
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    set_page_background(section)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("NUMERIQU")
    r.bold = True
    r.font.name = "Aptos Display"
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(20, 33, 61)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Engineering Onboarding Handbook")
    r.font.name = "Aptos Display"
    r.font.size = Pt(22)
    r.bold = True
    r.font.color.rgb = RGBColor(36, 71, 120)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "A single-source onboarding document for new engineers joining the Numeriqu platform."
    )
    r.font.size = Pt(11)
    r.italic = True

    for _ in range(4):
        document.add_paragraph("")

    intro = document.add_paragraph(style="Callout")
    intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    intro.add_run(
        "This handbook is written in the style used by strong software organizations: "
        "it explains not only what the repository contains, but why the product exists, "
        "how the platform is shaped, what tradeoffs guided its design, and how a new engineer "
        "is expected to work from day one."
    )

    table = document.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    meta = [
        ("Document Type", "Engineering onboarding and architecture handbook"),
        ("Audience", "New software engineers, technical leads, delivery partners"),
        ("Primary Scope", "Product context, repository structure, platform architecture, engineering workflow"),
        ("Repository", "Numeriqu monorepo"),
    ]
    for row_idx, (left, right) in enumerate(meta):
        row = table.rows[row_idx]
        row.cells[0].text = left
        row.cells[1].text = right
        set_cell_shading(row.cells[0], "EEF3F8")

    document.add_paragraph("")
    quote = document.add_paragraph(style="Callout")
    quote.add_run(
        "Guiding idea: a new engineer should be able to read this document, understand the story of the business, "
        "the architecture of the platform, the shape of the codebase, and the expectations of the team before writing their first line of production code."
    )

    document.add_section(WD_SECTION.NEW_PAGE)


def add_heading_paragraph(document, title, text):
    document.add_heading(title, level=1)
    document.add_paragraph(text)


def add_bullets(document, items):
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered(document, items):
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_footer(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Numeriqu Engineering Onboarding Handbook | ")
    add_page_number(footer)


def build_document():
    document = Document()
    configure_styles(document)
    add_cover(document)

    add_heading_paragraph(
        document,
        "1. Executive Welcome",
        "Welcome to Numeriqu. This platform is being built around a simple but important belief: modern finance teams should not have to choose between operational systems, analytical clarity, and AI-assisted decision support. In many companies, financial truth is fragmented across ERP systems, manually maintained spreadsheets, dashboard tools, email threads, and executive reporting decks. Numeriqu exists to reduce that fragmentation and turn disconnected financial signals into a coherent operating system for organizations."
    )
    document.add_paragraph(
        "As a new engineer, your role is not limited to implementing isolated tickets. You are joining a product effort that aims to combine data plumbing, analytics, secure multi-tenant application design, and AI-assisted workflows into a trustworthy experience for organizations. That means the engineering standard here is not just correctness in code, but clarity in systems thinking, respect for data boundaries, and care for maintainability."
    )

    add_heading_paragraph(
        document,
        "2. The Numeriqu Story",
        "The easiest way to understand Numeriqu is to begin with the problem it addresses. In a typical finance environment, raw accounting data is captured in systems such as QuickBooks, Xero, Workday, or Dynamics 365. Those systems are essential for recording transactions, but they are rarely the best place to answer strategic business questions. Teams export data, re-model it in spreadsheets or BI tools, and then spend additional time explaining the same numbers in meetings, status updates, and executive narratives."
    )
    document.add_paragraph(
        "Numeriqu compresses that gap. The product ingests source data, normalizes it into a shared analytical model, surfaces dashboards for routine visibility, and supports AI-driven workflows that can explain, summarize, and shape new dashboard views. In other words, the product is not just a reporting layer. It is a structured environment for turning financial data into decisions."
    )

    add_heading_paragraph(
        document,
        "3. Why Numeriqu Matters",
        "The strongest product organizations are explicit about why their systems exist. Numeriqu matters because it is designed to solve a business coordination problem, not only a technical one. Financial truth is frequently delayed, duplicated, or misinterpreted when operational systems, analysis layers, and decision-making conversations are split apart."
    )
    add_bullets(
        document,
        [
            "It gives organizations a shared, organization-scoped workspace rather than scattered single-purpose tools.",
            "It separates transactional truth from analytical truth so that performance reporting does not overload operational systems.",
            "It uses AI only where it can be grounded by real data and constrained by system design.",
            "It treats security, authorization, and tenant isolation as foundational architecture rather than afterthoughts.",
            "It creates a path from raw finance data to dashboards, chat, and action-oriented interpretation inside one product surface.",
        ],
    )

    add_heading_paragraph(
        document,
        "4. Product Overview",
        "Numeriqu is best understood as a set of coordinated capabilities rather than a single screen or workflow. Each capability is intentionally scoped, but they are designed to reinforce each other."
    )
    add_bullets(
        document,
        [
            "Authentication and access management: email-based OTP flows, organization memberships, invitations, and permissions.",
            "Integrations: provider connectivity and sync orchestration for financial source systems.",
            "Analytics: KPI and trend views backed by transformed analytical data.",
            "Dashboards: persisted dashboard definitions, widget metadata, and sharing.",
            "Messaging and collaboration: conversation flows linked to organizational context.",
            "RAG experiences: grounded finance chat with contextual data retrieval.",
            "Agent experiences: question-to-dashboard or question-to-chart planning with controlled execution paths.",
        ],
    )

    add_heading_paragraph(
        document,
        "5. Who the Product Is For",
        "Although the repository is technical, product framing matters because it shapes architecture decisions. Numeriqu is aimed at organization-centered workflows rather than individual note-taking or ad hoc query tools. Its implied users include finance leaders, operators, analysts, and teams who need shared visibility into financial performance."
    )
    document.add_paragraph(
        "This user model explains many design choices in the codebase. Organization membership is a first-class concept. Shared dashboards matter. Permissions matter. Auditability matters. Data freshness matters. The platform is designed for environments where multiple people need to trust, interpret, and act on the same financial context."
    )

    add_heading_paragraph(
        document,
        "6. Engineering Principles Behind the Product",
        "Top engineering organizations document not only what exists, but the principles that should remain stable as the system evolves. The following principles are visible throughout the repository and should guide future changes."
    )
    add_bullets(
        document,
        [
            "Organization-first design: most meaningful business objects are scoped to an organization.",
            "Separation of concerns: frontend, API, transactional storage, and analytical storage each have a defined role.",
            "Grounded AI: AI features are useful only when constrained by trustworthy data access patterns.",
            "Operational traceability: actions that affect dashboards, messages, or agent behavior should be auditable.",
            "Maintainable modularity: the backend is organized by domain responsibility rather than as one large service blob.",
        ],
    )

    add_heading_paragraph(
        document,
        "7. Repository Overview",
        "The Numeriqu repository is a monorepo managed with pnpm workspaces and Turborepo. This choice allows the team to keep the product surface, API, database package, analytics models, and shared configuration in one place while still preserving strong internal boundaries."
    )
    add_bullets(
        document,
        [
            "apps/web: the main Next.js frontend application.",
            "apps/api: the main NestJS backend service.",
            "apps/docs: a documentation app shell that currently is not the primary engineering source of truth.",
            "packages/db: Prisma schema, migrations, generated client, and data scripts.",
            "packages/analytics: dbt transformation layer for ClickHouse analytics models.",
            "packages/ui: shared UI components.",
            "docs: engineering documentation and architecture notes.",
        ],
    )

    add_heading_paragraph(
        document,
        "8. The Architecture in Plain Language",
        "A good onboarding document should let an engineer explain the architecture in one conversation without opening the code. Numeriqu has four major technical layers. The web application delivers the user experience. The API owns business logic, authentication orchestration, organization scoping, and AI coordination. Postgres stores transactional system-of-record entities. ClickHouse stores analytical facts and transformed marts used by dashboards and AI grounding. Around these layers sit supporting services such as Redis, Resend, Supabase Auth, and the chosen LLM provider."
    )

    add_heading_paragraph(
        document,
        "9. Frontend Architecture",
        "The frontend is implemented in Next.js with the App Router model. It is responsible for the user journey from authentication through dashboard usage, integrations, messaging, and AI workspace flows. The frontend is not the source of authority for security-sensitive behavior. It is a consumer of backend-controlled authentication, organization scoping, and permission enforcement."
    )
    add_bullets(
        document,
        [
            "The landing and marketing surfaces communicate the product narrative.",
            "The dashboard workspace hosts authenticated product flows.",
            "The frontend API wrappers centralize transport behavior and typed contracts.",
            "Cookie-aware backend communication is part of the intended runtime design.",
        ],
    )

    add_heading_paragraph(
        document,
        "10. Backend Architecture",
        "The backend is implemented in NestJS and composed through domain modules. This is an intentional architecture choice. Rather than centralizing every concern into a single service, the application root imports modules for auth, organization context, integrations, analytics, dashboards, messaging, RAG, agent, audit, health, and related supporting concerns. This structure scales better for complex product logic because it keeps responsibilities visible and reduces accidental coupling."
    )
    document.add_paragraph(
        "For a new engineer, one of the most important mental models is that the backend is the policy and orchestration boundary. It is where authorization is enforced, where organization context is resolved, where data movement decisions are coordinated, and where AI behavior is tied back to the rest of the system."
    )

    add_heading_paragraph(
        document,
        "11. Transactional Data Model",
        "Postgres, modeled through Prisma in packages/db, stores the transactional state of the product. This includes users, organizations, memberships, invites, ERP connections, sync jobs, dashboard metadata, messaging records, and AI session metadata. These are the entities that represent who is using the product, what organization they belong to, what connections exist, and what collaboration or workflow artifacts have been created."
    )
    document.add_paragraph(
        "The practical rule is simple: if a record represents an operational or business workflow object, it likely belongs in Postgres."
    )

    add_heading_paragraph(
        document,
        "12. Analytical Data Model",
        "ClickHouse is the analytical read store. It is where finance facts are queried at speed and where the application reads from transformed data structures instead of repeatedly re-deriving business meaning from raw transactional sources. The dbt project in packages/analytics creates the modeling discipline around this layer, using provider-specific staging models and shared Gold-layer marts."
    )
    add_bullets(
        document,
        [
            "Bronze or raw data preserves source-shaped ingested records.",
            "Silver or staging models clean and normalize provider-specific structures.",
            "Gold or marts present stable, product-facing analytical models for dashboards and AI flows.",
        ],
    )

    add_heading_paragraph(
        document,
        "13. Authentication and Identity",
        "Numeriqu uses a backend-owned authentication model. Supabase provides identity infrastructure and session mechanics, but the application logic around OTP flows, verification, session issuance, and organization-related behavior lives in the API. Redis supports short-lived OTP state and throttling. Resend handles email delivery. The frontend does not directly own the critical trust decisions in the authentication flow."
    )
    document.add_paragraph(
        "This is an important design decision because it centralizes security-sensitive logic and keeps the product’s business rules around sign-in, invite acceptance, and organization alignment in one place."
    )

    add_heading_paragraph(
        document,
        "14. Organization and Permission Model",
        "Numeriqu is deliberately organization-first. A user does not merely exist; they exist in relation to organizations and memberships. Permissions are not decorative metadata. They shape which dashboards can be viewed, created, or shared, and they determine whether a user can act within a given organizational context."
    )
    add_bullets(
        document,
        [
            "Organization membership must be validated for scoped operations.",
            "Permission checks must be enforced at the backend boundary.",
            "Queries should not assume a user’s default organization when an explicit organization scope is required.",
            "Every engineer should treat tenant isolation as a structural invariant.",
        ],
    )

    add_heading_paragraph(
        document,
        "15. Integrations and Data Flow",
        "Integrations are the bridge between source systems and Numeriqu’s internal models. Provider-specific logic exists for systems such as QuickBooks, Xero, Workday, and Dynamics 365. Sync orchestration, connection metadata, and job lifecycle records are managed in the application layer. The analytical side of the platform depends on this ingestion pipeline being both observable and disciplined."
    )
    document.add_paragraph(
        "From a business perspective, integrations allow Numeriqu to speak to the systems where financial truth originates. From an engineering perspective, integrations are a reliability boundary and should be treated with the same seriousness as auth or billing in other product categories."
    )

    add_heading_paragraph(
        document,
        "16. Dashboards, Messaging, and Collaboration",
        "Dashboards are not static images in Numeriqu. They are application-managed artifacts with ownership, widgets, sharing rules, and lifecycle updates. Messaging is similarly organization-scoped and durable. These capabilities matter because Numeriqu is designed to support collaborative understanding, not just solo exploration."
    )
    document.add_paragraph(
        "A recurring theme in the codebase is that collaborative artifacts should remain auditable, permission-aware, and tied back to the right organization. That is exactly the kind of design discipline strong companies encode into their product foundation early."
    )

    add_heading_paragraph(
        document,
        "17. AI, RAG, and Agent Experiences",
        "Numeriqu includes two distinct AI-oriented patterns. The first is retrieval-oriented chat, where the user asks questions and the system responds using grounded context. The second is agent-style orchestration, where a question may drive structured planning, chart selection, query execution, or dashboard creation. These systems are not simply large prompts wrapped in APIs. They are application features that must respect permissions, data grounding, and runtime observability."
    )
    document.add_paragraph(
        "One of the healthiest engineering instincts in this repository is the recognition that AI should be constrained by architecture rather than trusted by default. New engineers should preserve that instinct."
    )

    add_heading_paragraph(
        document,
        "18. What a New Engineer Should Understand in the First Week",
        "A strong onboarding process makes expectations explicit. By the end of your first week, you should be able to explain the product story, navigate the monorepo confidently, identify the role of each major app or package, describe the difference between Postgres and ClickHouse responsibilities, and trace at least one end-to-end product flow across frontend, backend, and data layers."
    )
    add_numbered(
        document,
        [
            "Read the onboarding, repository structure, architecture, and workflow documents.",
            "Run the application locally and confirm the web and API surfaces start successfully.",
            "Inspect the Prisma schema to understand organization-centric modeling.",
            "Read the agent architecture note before changing agent-adjacent code.",
            "Trace one user flow such as login, dashboard access, or an AI question path.",
        ],
    )

    add_heading_paragraph(
        document,
        "19. Expected Engineering Behavior",
        "Professional product teams rely on clear engineering behavior, not only technical skill. In Numeriqu, engineers are expected to work with architectural awareness. That means understanding where code belongs, being careful with organization scoping, preserving security-sensitive boundaries, and avoiding quick fixes that create hidden coupling."
    )
    add_bullets(
        document,
        [
            "Prefer clarity over cleverness.",
            "Respect system boundaries before optimizing local convenience.",
            "Update documentation when you materially change architecture or workflow.",
            "Validate permission and tenant effects whenever a feature touches shared data.",
            "Use the narrowest effective tests first, then broader validation when the change crosses boundaries.",
        ],
    )

    add_heading_paragraph(
        document,
        "20. Development Workflow",
        "The repository is operated through common root commands and targeted package commands. From the root, the essential workflow is to install dependencies, configure environment variables, run the workspace, and validate changes through linting, type checking, builds, and targeted tests. Because the system spans multiple technical layers, the right validation strategy depends on the type of change being made."
    )
    add_bullets(
        document,
        [
            "Root commands: pnpm install, pnpm dev, pnpm build, pnpm lint, pnpm check-types.",
            "Backend-focused work: targeted API build and test commands.",
            "Database-focused work: Prisma generate or migration commands in the db package.",
            "Agent-focused work: grounding, parity, and browser-run QA where relevant.",
        ],
    )

    add_heading_paragraph(
        document,
        "21. How to Think About Risk in This Codebase",
        "Not all files carry the same risk. UI-only presentation changes are usually simpler than changes to auth, organization context, integration syncing, or AI data grounding. A mature engineering culture makes that distinction visible so engineers know where extra care is required."
    )
    add_bullets(
        document,
        [
            "High-risk zones: auth, org context, integration orchestration, schema changes, AI grounding paths.",
            "Medium-risk zones: dashboards, sharing, messaging persistence, frontend API state management.",
            "Lower-risk zones: isolated UI improvements and documentation updates.",
        ],
    )

    add_heading_paragraph(
        document,
        "22. Recommended Reading Order Inside the Repository",
        "A new engineer should not start by opening random files. The most efficient order is to begin with the main README, then the docs hub, then the onboarding and repository structure documents, then the system architecture and schema documents, and finally the domain-specific architecture notes for the subsystem you are about to change."
    )
    add_bullets(
        document,
        [
            "README.md",
            "docs/README.md",
            "docs/developer-onboarding.md",
            "docs/repository-structure.md",
            "docs/architecture.md",
            "docs/database-schema-numeriqu.md",
            "apps/api/src/modules/agent/AGENT_ARCHITECTURE.md for agent-related work",
        ],
    )

    add_heading_paragraph(
        document,
        "23. What Good Looks Like Here",
        "Good engineering work in Numeriqu is work that leaves the system clearer, not merely changed. A good contribution respects tenant boundaries, fits the existing architecture, keeps the distinction between transactional and analytical concerns intact, and makes the next engineer’s job easier. A good document does the same thing for understanding."
    )

    add_heading_paragraph(
        document,
        "24. Closing Note",
        "The best onboarding documents do not attempt to impress people with volume. They reduce ambiguity, communicate intent, and help a new team member feel that the system has been thought through. That is the purpose of this handbook. It should serve as a high-trust starting point for anyone joining Numeriqu engineering and should evolve with the product as the architecture deepens."
    )

    appendix = document.add_heading("Appendix A. Quick Reference", level=1)
    appendix.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_bullets(
        document,
        [
            "Primary frontend: apps/web",
            "Primary backend: apps/api",
            "Transactional schema: packages/db/prisma/schema.prisma",
            "Analytics layer: packages/analytics",
            "Architecture docs: docs/",
            "Default local URLs: web on localhost:3001, api on localhost:3000",
        ],
    )

    for section in document.sections:
        add_footer(section)

    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
